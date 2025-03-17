import tkinter as tk
from tkinter import filedialog, messagebox
import os
from docx import Document

def list_tables(doc_path):
    """Extracts tables from the document and shows their index with a preview."""
    try:
        doc = Document(doc_path)
        table_listbox.delete(0, tk.END) # Clear previous list

        if not doc.tables:
            messagebox.showinfo("Info", "No tables found in this document.")
            return

        for i, table in enumerate(doc.tables):
            # Preview: Show first row's content
            preview = " | ".join([cell.text.strip() for cell in table.rows[0].cells]) if table.rows else "(Empty Table)"
            table_listbox.insert(tk.END, f"Table {i}: {preview}")

    except Exception as e:
        messagebox.showerror("Error", f"Failed to read tables: {str(e)}")

def select_table():
    """Gets the selected table index from the listbox."""
    try:
        selected_index = table_listbox.curselection()[0]
        index_entry.delete(0, tk.END)
        index_entry.insert(0, selected_index)
    except IndexError:
        messagebox.showerror("Error", "Please select a table from the list.")

def replace_table_in_file(doc_path, new_table_path, table_index):
    """Replaces a table at a given index while retaining formatting in a single file."""
    try:
        doc = Document(doc_path)
        new_table_doc = Document(new_table_path)

        if table_index >= len(doc.tables):
            return f"Skipped {os.path.basename(doc_path)} (Table index {table_index} not found)"

        old_table = doc.tables[table_index]
        new_table = new_table_doc.tables[0] if new_table_doc.tables else None

        if not new_table:
            return f"Skipped {os.path.basename(doc_path)} (No table found in replacement doc)"

        # Preserve table position
        old_table_parent = old_table._element.getparent()
        new_table_element = new_table._element

        # Copy old table's style
        new_table_element.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblStyle'] = old_table._element.attrib.get(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblStyle', ''
        )

        # Replace the old table
        old_table_parent.insert(old_table_parent.index(old_table._element), new_table_element)
        old_table_parent.remove(old_table._element)

        doc.save(doc_path)
        return f"Processed {os.path.basename(doc_path)} successfully"

    except Exception as e:
        return f"Error processing {os.path.basename(doc_path)}: {str(e)}"

def select_folder():
    """Allows the user to select a folder containing multiple Word documents."""
    folder_path = filedialog.askdirectory()
    folder_entry.delete(0, tk.END)
    folder_entry.insert(0, folder_path)

def select_new_table():
    """Allows the user to select a Word document containing the new table."""
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    new_table_entry.delete(0, tk.END)
    new_table_entry.insert(0, file_path)

def run_batch_replacement():
    """Replaces the table in all Word documents in the selected folder."""
    folder_path = folder_entry.get()
    new_table_path = new_table_entry.get()
    
    if not os.path.isdir(folder_path):
        messagebox.showerror("Error", "Please select a valid folder.")
        return
    
    try:
        table_index = int(index_entry.get())
    except ValueError:
        messagebox.showerror("Error", "Please select a valid table index.")
        return
    
    results = []
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            result = replace_table_in_file(file_path, new_table_path, table_index)
            results.append(result)

    messagebox.showinfo("Batch Processing Complete", "\n".join(results))

# GUI Window
root = tk.Tk()
root.title("Batch Word Table Replacer")

tk.Label(root, text="Select Folder with Word Documents:").grid(row=0, column=0, padx=10, pady=5)
folder_entry = tk.Entry(root, width=50)
folder_entry.grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=select_folder).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="Select New Table Document:").grid(row=1, column=0, padx=10, pady=5)
new_table_entry = tk.Entry(root, width=50)
new_table_entry.grid(row=1, column=1, padx=10, pady=5)
tk.Button(root, text="Browse", command=select_new_table).grid(row=1, column=2, padx=10, pady=5)

tk.Label(root, text="Select Table to Replace:").grid(row=2, column=0, padx=10, pady=5)
table_listbox = tk.Listbox(root, height=6, width=70)
table_listbox.grid(row=2, column=1, columnspan=2, padx=10, pady=5)
tk.Button(root, text="Select", command=select_table).grid(row=3, column=2, padx=10, pady=5)

tk.Label(root, text="Table Index:").grid(row=4, column=0, padx=10, pady=5)
index_entry = tk.Entry(root, width=10)
index_entry.grid(row=4, column=1, padx=10, pady=5)

tk.Button(root, text="Replace Table in All Files", command=run_batch_replacement).grid(row=5, column=1, pady=10)

root.mainloop()

